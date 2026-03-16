"""
CGS — Compliance report generator (multi-framework).

Supported frameworks:
  - NIS2 Directive (EU 2022/2555)
  - ISO 27001:2022 Annex A
  - SOC 2 Type II (Trust Service Criteria)
  - NIST CSF 2.0 (Cybersecurity Framework)
  - CIS Controls v8
  - GDPR Article 32 (Data protection)
  - DORA (Digital Operational Resilience Act — EU financial sector)
  - Cyber Essentials (UK)

Users and admins can select which frameworks to include.
All are enabled by default.
"""

import io
import logging
import os
from datetime import datetime, timedelta

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak,
)

logger = logging.getLogger("cgs.compliance")

# ── Colors ──
C_DARK = HexColor("#0F172A")
C_BLUE = HexColor("#1E40AF")
C_LIGHT_BLUE = HexColor("#3B82F6")
C_GREEN = HexColor("#16A34A")
C_YELLOW = HexColor("#CA8A04")
C_ORANGE = HexColor("#EA580C")
C_RED = HexColor("#DC2626")
C_GRAY = HexColor("#64748B")
C_LIGHT_GRAY = HexColor("#F1F5F9")
C_BORDER = HexColor("#CBD5E1")
C_WHITE = white
C_BG_GREEN = HexColor("#F0FDF4")
C_BG_RED = HexColor("#FEF2F2")


def _styles():
    base = getSampleStyleSheet()
    s = {}
    s["cover_title"] = ParagraphStyle("CT", parent=base["Title"], fontSize=28,
        textColor=C_DARK, alignment=TA_CENTER, fontName="Helvetica-Bold",
        spaceAfter=4*mm, leading=34)
    s["cover_sub"] = ParagraphStyle("CS", parent=base["Normal"], fontSize=14,
        textColor=C_GRAY, alignment=TA_CENTER, spaceAfter=2*mm)
    s["h1"] = ParagraphStyle("H1", parent=base["Heading1"], fontSize=16,
        textColor=C_BLUE, fontName="Helvetica-Bold", spaceBefore=10*mm,
        spaceAfter=4*mm, leading=20)
    s["h2"] = ParagraphStyle("H2", parent=base["Heading2"], fontSize=13,
        textColor=C_DARK, fontName="Helvetica-Bold", spaceBefore=6*mm,
        spaceAfter=3*mm, leading=16)
    s["h3"] = ParagraphStyle("H3", parent=base["Heading3"], fontSize=11,
        textColor=C_BLUE, fontName="Helvetica-Bold", spaceBefore=4*mm,
        spaceAfter=2*mm)
    s["body"] = ParagraphStyle("B", parent=base["Normal"], fontSize=10,
        textColor=C_DARK, leading=14, alignment=TA_JUSTIFY, spaceAfter=2*mm)
    s["small"] = ParagraphStyle("SM", parent=base["Normal"], fontSize=8,
        textColor=C_GRAY, leading=10)
    s["label"] = ParagraphStyle("L", parent=base["Normal"], fontSize=9, textColor=C_GRAY)
    s["value"] = ParagraphStyle("V", parent=base["Normal"], fontSize=10,
        textColor=C_DARK, fontName="Helvetica-Bold")
    s["rec"] = ParagraphStyle("REC", parent=base["Normal"], fontSize=9,
        textColor=C_DARK, leading=13, spaceAfter=1*mm, leftIndent=10*mm)
    s["score_big"] = ParagraphStyle("SB", parent=base["Normal"], fontSize=48,
        textColor=C_DARK, fontName="Helvetica-Bold", alignment=TA_CENTER)
    s["score_label"] = ParagraphStyle("SL", parent=base["Normal"], fontSize=12,
        textColor=C_GRAY, alignment=TA_CENTER, spaceAfter=4*mm)
    s["footer"] = ParagraphStyle("F", parent=base["Normal"], fontSize=7,
        textColor=C_GRAY, alignment=TA_CENTER)
    s["pass"] = ParagraphStyle("PASS", parent=base["Normal"], fontSize=9,
        textColor=C_GREEN, fontName="Helvetica-Bold")
    s["fail"] = ParagraphStyle("FAIL", parent=base["Normal"], fontSize=9,
        textColor=C_RED, fontName="Helvetica-Bold")
    return s


# ══════════════════════════════════════════════════
# Framework definitions
# ══════════════════════════════════════════════════

FRAMEWORKS = {
    "nis2": {"name": "NIS2 Directive", "full": "EU Directive 2022/2555 on Network and Information Security", "version": "2022"},
    "iso27001": {"name": "ISO 27001", "full": "ISO/IEC 27001:2022 Information Security Management", "version": "2022"},
    "soc2": {"name": "SOC 2 Type II", "full": "AICPA Service Organization Control 2 — Trust Service Criteria", "version": "2017"},
    "nist_csf": {"name": "NIST CSF 2.0", "full": "NIST Cybersecurity Framework 2.0", "version": "2024"},
    "cis": {"name": "CIS Controls", "full": "CIS Critical Security Controls v8", "version": "v8"},
    "gdpr": {"name": "GDPR Art.32", "full": "EU General Data Protection Regulation — Article 32 Security of Processing", "version": "2016/679"},
    "dora": {"name": "DORA", "full": "EU Digital Operational Resilience Act (Regulation 2022/2554)", "version": "2022"},
    "cyber_essentials": {"name": "Cyber Essentials", "full": "UK Cyber Essentials Certification Scheme", "version": "2023"},
}

ALL_FRAMEWORK_IDS = list(FRAMEWORKS.keys())

# ══════════════════════════════════════════════════
# Controls with multi-framework mapping
# ══════════════════════════════════════════════════
# Each control: (id, category, title, description, check_fn, weight, recommendation, mappings)
# mappings = {framework_id: reference}

CONTROLS = [
    # ── Network monitoring ──
    ("NET-01", "Network", "Real-time packet capture",
     "Network traffic is captured and analyzed in real time",
     "check_sniffer_active", 10,
     "Enable and verify the packet sniffer is running on the correct interface",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.2", "nist_csf": "DE.CM-01",
      "cis": "CIS 13.1", "gdpr": "Art.32(1)(b)", "dora": "Art.10(1)", "cyber_essentials": "—"}),

    ("NET-02", "Network", "Network discovery",
     "All devices on the network are inventoried automatically",
     "check_host_inventory", 8,
     "Run an ARP discovery scan and verify all network segments are covered",
     {"nis2": "Art.21(2)(a)", "iso27001": "A.8.9", "soc2": "CC6.1", "nist_csf": "ID.AM-01",
      "cis": "CIS 1.1", "gdpr": "Art.32(1)(d)", "dora": "Art.8(1)", "cyber_essentials": "—"}),

    ("NET-03", "Network", "DNS monitoring",
     "DNS queries are logged and analyzed for tunneling/DGA",
     "check_dns_monitoring", 6,
     "Verify DNS logging is active and entropy analysis threshold is set",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.2", "nist_csf": "DE.CM-01",
      "cis": "CIS 13.6", "gdpr": "—", "dora": "Art.10(1)", "cyber_essentials": "—"}),

    # ── Threat detection ──
    ("DET-01", "Detection", "Intrusion detection active",
     "Threat engine analyzes traffic for known attack patterns",
     "check_threat_engine", 10,
     "Verify the threat engine is running and processing events",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.2", "nist_csf": "DE.AE-02",
      "cis": "CIS 13.3", "gdpr": "Art.32(1)(b)", "dora": "Art.10(1)", "cyber_essentials": "Malware protection"}),

    ("DET-02", "Detection", "Advanced detectors enabled",
     "Advanced behavioral detectors are active (not just observing)",
     "check_advanced_detectors", 6,
     "Switch advanced detectors from observation to active mode after baseline period",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.2", "nist_csf": "DE.AE-03",
      "cis": "CIS 13.3", "gdpr": "—", "dora": "Art.10(2)", "cyber_essentials": "—"}),

    ("DET-03", "Detection", "Threat intelligence feeds",
     "IOC feeds are configured and updated within last 24h",
     "check_threat_feeds", 8,
     "Verify threat feed auto-refresh is working. Check /api/threat-feeds for last update time",
     {"nis2": "Art.21(2)(e)", "iso27001": "A.5.7", "soc2": "CC7.1", "nist_csf": "ID.RA-02",
      "cis": "CIS 13.8", "gdpr": "—", "dora": "Art.13(1)", "cyber_essentials": "—"}),

    ("DET-04", "Detection", "Kill chain detection",
     "Multi-step attack sequence detection is active",
     "check_killchain", 4,
     "The kill chain detector should be enabled to correlate multi-step attacks",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.3", "nist_csf": "DE.AE-04",
      "cis": "CIS 13.5", "gdpr": "—", "dora": "Art.10(1)", "cyber_essentials": "—"}),

    # ── Active defense ──
    ("DEF-01", "Defense", "Firewall integration",
     "Active defense can block threats via iptables/nftables",
     "check_firewall_active", 10,
     "Install iptables or nftables and verify the CGS chain is created",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.8.23", "soc2": "CC6.6", "nist_csf": "PR.IR-01",
      "cis": "CIS 4.4", "gdpr": "Art.32(1)(b)", "dora": "Art.9(2)", "cyber_essentials": "Firewalls"}),

    ("DEF-02", "Defense", "Auto-block enabled",
     "Critical threats are automatically blocked without admin delay",
     "check_auto_block", 6,
     "Enable defense.auto_block in configuration for immediate response",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.8.23", "soc2": "CC7.4", "nist_csf": "RS.MI-01",
      "cis": "CIS 13.4", "gdpr": "—", "dora": "Art.10(5)", "cyber_essentials": "—"}),

    ("DEF-03", "Defense", "Graduated response",
     "Threat response uses escalation levels (monitor to block)",
     "check_escalation", 4,
     "The escalation ladder provides proportional response",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.8.23", "soc2": "CC7.4", "nist_csf": "RS.MI-02",
      "cis": "CIS 13.4", "gdpr": "—", "dora": "Art.10(5)", "cyber_essentials": "—"}),

    # ── Incident response ──
    ("INC-01", "Incident Response", "Incident workflow configured",
     "Incident detection triggers admin notification within defined timeout",
     "check_incident_workflow", 8,
     "Configure email notifications and set approval timeout under email settings",
     {"nis2": "Art.21(2)(d)", "iso27001": "A.5.24", "soc2": "CC7.3", "nist_csf": "RS.CO-02",
      "cis": "CIS 17.2", "gdpr": "Art.33", "dora": "Art.17(1)", "cyber_essentials": "—"}),

    ("INC-02", "Incident Response", "Forensic collection",
     "Forensic evidence is automatically collected after incidents",
     "check_forensic_enabled", 6,
     "Enable client_agent.collect_after_incident and verify forensic reports are generated",
     {"nis2": "Art.21(2)(d)", "iso27001": "A.5.28", "soc2": "CC7.5", "nist_csf": "RS.AN-03",
      "cis": "CIS 17.6", "gdpr": "—", "dora": "Art.17(3)", "cyber_essentials": "—"}),

    ("INC-03", "Incident Response", "Admin notification",
     "At least one admin email is configured for alerts",
     "check_admin_emails", 8,
     "Add at least one admin email address in the email configuration section",
     {"nis2": "Art.23", "iso27001": "A.5.25", "soc2": "CC2.3", "nist_csf": "RS.CO-03",
      "cis": "CIS 17.3", "gdpr": "Art.33", "dora": "Art.19(1)", "cyber_essentials": "—"}),

    # ── Access control ──
    ("ACC-01", "Access Control", "Web dashboard authentication",
     "Dashboard requires authentication with strong passwords",
     "check_auth_configured", 8,
     "Ensure all accounts have passwords of 16+ characters",
     {"nis2": "Art.21(2)(i)", "iso27001": "A.8.5", "soc2": "CC6.1", "nist_csf": "PR.AA-01",
      "cis": "CIS 5.2", "gdpr": "Art.32(1)(b)", "dora": "Art.9(4)(c)", "cyber_essentials": "Access control"}),

    ("ACC-02", "Access Control", "Two-factor authentication",
     "2FA (TOTP) is enabled for at least admin accounts",
     "check_2fa", 6,
     "Configure TOTP for admin accounts via Users settings",
     {"nis2": "Art.21(2)(i)", "iso27001": "A.8.5", "soc2": "CC6.1", "nist_csf": "PR.AA-03",
      "cis": "CIS 6.3", "gdpr": "Art.32(1)(b)", "dora": "Art.9(4)(c)", "cyber_essentials": "Access control"}),

    ("ACC-03", "Access Control", "Login protection",
     "Login lockout is configured against brute force",
     "check_login_lockout", 4,
     "Verify web.max_login_attempts and web.lockout_duration_seconds are set",
     {"nis2": "Art.21(2)(i)", "iso27001": "A.8.5", "soc2": "CC6.1", "nist_csf": "PR.AA-01",
      "cis": "CIS 5.6", "gdpr": "—", "dora": "Art.9(4)(c)", "cyber_essentials": "Access control"}),

    # ── System hardening ──
    ("SYS-01", "Hardening", "TLS encryption",
     "Web interface uses HTTPS with TLS certificate",
     "check_tls", 8,
     "Generate or install a TLS certificate for the web dashboard",
     {"nis2": "Art.21(2)(h)", "iso27001": "A.8.24", "soc2": "CC6.7", "nist_csf": "PR.DS-02",
      "cis": "CIS 3.10", "gdpr": "Art.32(1)(a)", "dora": "Art.9(3)(b)", "cyber_essentials": "Secure configuration"}),

    ("SYS-02", "Hardening", "SSH hardened",
     "SSH password auth disabled when keys exist, root login restricted",
     "check_ssh", 6,
     "Run SSH hardening to disable password authentication when SSH keys are configured",
     {"nis2": "Art.21(2)(h)", "iso27001": "A.8.9", "soc2": "CC6.1", "nist_csf": "PR.IR-01",
      "cis": "CIS 4.1", "gdpr": "—", "dora": "Art.9(2)", "cyber_essentials": "Secure configuration"}),

    ("SYS-03", "Hardening", "Code integrity",
     "Application files are verified against integrity manifest",
     "check_integrity", 6,
     "Regenerate the integrity manifest and investigate any modified files",
     {"nis2": "Art.21(2)(e)", "iso27001": "A.8.9", "soc2": "CC8.1", "nist_csf": "PR.DS-06",
      "cis": "CIS 2.7", "gdpr": "—", "dora": "Art.9(1)", "cyber_essentials": "—"}),

    ("SYS-04", "Hardening", "Backup configured",
     "Automatic backups are configured and running",
     "check_backup", 6,
     "Configure backup.directory and verify daily backup job runs at 04:00",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.8.13", "soc2": "A1.2", "nist_csf": "PR.IR-04",
      "cis": "CIS 11.2", "gdpr": "Art.32(1)(c)", "dora": "Art.12(1)", "cyber_essentials": "—"}),

    # ── Logging & audit ──
    ("LOG-01", "Audit", "Audit trail integrity",
     "Tamper-proof audit chain is active and verified",
     "check_audit_chain", 8,
     "The hash-chain audit log should pass integrity verification",
     {"nis2": "Art.21(2)(g)", "iso27001": "A.8.15", "soc2": "CC7.2", "nist_csf": "DE.AE-06",
      "cis": "CIS 8.2", "gdpr": "Art.32(1)(d)", "dora": "Art.12(3)", "cyber_essentials": "—"}),

    ("LOG-02", "Audit", "Data retention policy",
     "Alert and log retention periods are configured",
     "check_retention", 4,
     "Set appropriate retention periods for alerts, events, and flows",
     {"nis2": "Art.21(2)(g)", "iso27001": "A.8.10", "soc2": "CC7.2", "nist_csf": "DE.AE-06",
      "cis": "CIS 8.1", "gdpr": "Art.5(1)(e)", "dora": "Art.12(3)", "cyber_essentials": "—"}),

    # ── Honeypot ──
    ("HON-01", "Deception", "Honeypot decoys active",
     "Decoy services detect reconnaissance with zero false positives",
     "check_honeypot", 4,
     "Enable honeypot with unused ports to detect attacker reconnaissance",
     {"nis2": "Art.21(2)(b)", "iso27001": "A.8.16", "soc2": "CC7.2", "nist_csf": "DE.CM-01",
      "cis": "CIS 13.1", "gdpr": "—", "dora": "Art.10(1)", "cyber_essentials": "—"}),
]

# ══════════════════════════════════════════════════
# Declarative controls (self-assessed by users/admins)
# ══════════════════════════════════════════════════
# (id, category, title, description, weight, prompt, mappings)

DECLARATIVE_CONTROLS = [
    # ── Governance & Policy ──
    ("ORG-01", "Governance", "Written security policy",
     "The organization has a documented and approved information security policy",
     8, "Is there a written and approved information security policy?",
     {"nis2": "Art.21(1)", "iso27001": "A.5.1", "soc2": "CC1.1", "nist_csf": "GV.PO-01",
      "cis": "CIS 1", "gdpr": "Art.24", "dora": "Art.5(1)", "cyber_essentials": "—"}),

    ("ORG-02", "Governance", "Security roles and responsibilities",
     "Security roles and responsibilities are defined and assigned",
     6, "Are security roles and responsibilities formally defined?",
     {"nis2": "Art.21(1)", "iso27001": "A.5.2", "soc2": "CC1.3", "nist_csf": "GV.RR-01",
      "cis": "CIS 1", "gdpr": "Art.24", "dora": "Art.5(2)(a)", "cyber_essentials": "—"}),

    ("ORG-03", "Governance", "Management commitment",
     "Senior management demonstrates commitment to information security",
     4, "Does senior management actively support and review the security program?",
     {"nis2": "Art.20(1)", "iso27001": "5.1", "soc2": "CC1.2", "nist_csf": "GV.RR-02",
      "cis": "—", "gdpr": "Art.24", "dora": "Art.5(2)(e)", "cyber_essentials": "—"}),

    ("ORG-04", "Governance", "Risk assessment process",
     "Regular risk assessments are conducted and documented",
     8, "Is there a formal risk assessment process performed at least annually?",
     {"nis2": "Art.21(2)(a)", "iso27001": "A.5.12", "soc2": "CC3.2", "nist_csf": "ID.RA-03",
      "cis": "CIS 1", "gdpr": "Art.32(1)", "dora": "Art.6(1)", "cyber_essentials": "—"}),

    # ── Human resources ──
    ("HR-01", "Human Resources", "Security awareness training",
     "All employees receive regular cybersecurity awareness training",
     8, "Do all employees receive cybersecurity training at least annually?",
     {"nis2": "Art.21(2)(g)", "iso27001": "A.6.3", "soc2": "CC1.4", "nist_csf": "PR.AT-01",
      "cis": "CIS 14.1", "gdpr": "Art.39(1)(b)", "dora": "Art.13(6)", "cyber_essentials": "—"}),

    ("HR-02", "Human Resources", "Phishing awareness",
     "Employees are trained to recognize and report phishing attempts",
     6, "Is there a phishing awareness program (simulation, training)?",
     {"nis2": "Art.21(2)(g)", "iso27001": "A.6.3", "soc2": "CC1.4", "nist_csf": "PR.AT-01",
      "cis": "CIS 14.2", "gdpr": "—", "dora": "Art.13(6)", "cyber_essentials": "—"}),

    ("HR-03", "Human Resources", "Background checks",
     "Background verification is performed for personnel in sensitive positions",
     4, "Are background checks performed for employees in security-sensitive roles?",
     {"nis2": "—", "iso27001": "A.6.1", "soc2": "CC1.4", "nist_csf": "PR.AT-02",
      "cis": "—", "gdpr": "—", "dora": "Art.5(2)(a)", "cyber_essentials": "—"}),

    # ── Physical security ──
    ("PHY-01", "Physical Security", "Physical access control",
     "Server rooms and network equipment are physically secured",
     6, "Are server rooms and network infrastructure physically access-controlled?",
     {"nis2": "Art.21(2)(i)", "iso27001": "A.7.1", "soc2": "CC6.4", "nist_csf": "PR.AA-06",
      "cis": "—", "gdpr": "Art.32(1)(b)", "dora": "Art.9(2)", "cyber_essentials": "—"}),

    ("PHY-02", "Physical Security", "Environmental protection",
     "Protection against fire, flood, and power failure is in place",
     4, "Is there protection against environmental threats (fire, flood, power failure)?",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.7.5", "soc2": "A1.1", "nist_csf": "PR.IR-04",
      "cis": "—", "gdpr": "—", "dora": "Art.12(1)", "cyber_essentials": "—"}),

    # ── Supply chain ──
    ("SUP-01", "Supply Chain", "Supplier security assessment",
     "Third-party suppliers are assessed for security risks",
     6, "Are suppliers and third-party providers assessed for security before engagement?",
     {"nis2": "Art.21(2)(d)", "iso27001": "A.5.19", "soc2": "CC9.2", "nist_csf": "GV.SC-03",
      "cis": "CIS 15.1", "gdpr": "Art.28", "dora": "Art.28(1)", "cyber_essentials": "—"}),

    ("SUP-02", "Supply Chain", "Supplier contracts include security",
     "Contracts with suppliers include information security requirements",
     4, "Do contracts with ICT suppliers include security clauses and audit rights?",
     {"nis2": "Art.21(2)(d)", "iso27001": "A.5.20", "soc2": "CC9.2", "nist_csf": "GV.SC-05",
      "cis": "CIS 15.2", "gdpr": "Art.28(3)", "dora": "Art.30(2)", "cyber_essentials": "—"}),

    # ── Business continuity ──
    ("BCP-01", "Business Continuity", "Business continuity plan",
     "A business continuity plan exists and is documented",
     8, "Is there a documented business continuity / disaster recovery plan?",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.5.30", "soc2": "A1.2", "nist_csf": "RC.RP-01",
      "cis": "CIS 11.1", "gdpr": "Art.32(1)(c)", "dora": "Art.11(1)", "cyber_essentials": "—"}),

    ("BCP-02", "Business Continuity", "BCP tested",
     "The business continuity plan has been tested within the last 12 months",
     6, "Has the BCP/DR plan been tested in the last 12 months?",
     {"nis2": "Art.21(2)(c)", "iso27001": "A.5.30", "soc2": "A1.3", "nist_csf": "RC.RP-03",
      "cis": "CIS 11.3", "gdpr": "Art.32(1)(d)", "dora": "Art.11(6)", "cyber_essentials": "—"}),

    # ── Data protection ──
    ("DAT-01", "Data Protection", "Data classification",
     "Data is classified by sensitivity level",
     6, "Is there a data classification scheme (public, internal, confidential, restricted)?",
     {"nis2": "—", "iso27001": "A.5.12", "soc2": "CC6.1", "nist_csf": "ID.AM-07",
      "cis": "CIS 3.1", "gdpr": "Art.5(1)(f)", "dora": "Art.9(4)(a)", "cyber_essentials": "—"}),

    ("DAT-02", "Data Protection", "Data Protection Officer",
     "A DPO or privacy officer is appointed (if required)",
     4, "Is a Data Protection Officer (DPO) or privacy officer appointed?",
     {"nis2": "—", "iso27001": "—", "soc2": "—", "nist_csf": "GV.RR-01",
      "cis": "—", "gdpr": "Art.37", "dora": "—", "cyber_essentials": "—"}),

    ("DAT-03", "Data Protection", "Privacy impact assessment",
     "Privacy impact assessments are conducted for high-risk processing",
     4, "Are privacy/data protection impact assessments performed for new systems?",
     {"nis2": "—", "iso27001": "—", "soc2": "P1.1", "nist_csf": "GV.PO-02",
      "cis": "—", "gdpr": "Art.35", "dora": "—", "cyber_essentials": "—"}),

    # ── Change management ──
    ("CHG-01", "Change Management", "Change management process",
     "A formal change management process exists for IT systems",
     6, "Is there a formal change management process for IT systems and configurations?",
     {"nis2": "Art.21(2)(e)", "iso27001": "A.8.32", "soc2": "CC8.1", "nist_csf": "PR.IP-03",
      "cis": "CIS 2.1", "gdpr": "—", "dora": "Art.9(4)(e)", "cyber_essentials": "Secure configuration"}),

    # ── Vulnerability management ──
    ("VUL-01", "Vulnerability Management", "Vulnerability scanning",
     "Regular vulnerability scans are performed on systems and applications",
     6, "Are vulnerability scans performed at least monthly on critical systems?",
     {"nis2": "Art.21(2)(e)", "iso27001": "A.8.8", "soc2": "CC7.1", "nist_csf": "ID.RA-01",
      "cis": "CIS 7.1", "gdpr": "Art.32(1)(d)", "dora": "Art.9(3)(a)", "cyber_essentials": "Security update management"}),

    ("VUL-02", "Vulnerability Management", "Patch management",
     "Security patches are applied within defined timeframes",
     8, "Are critical security patches applied within 14 days of release?",
     {"nis2": "Art.21(2)(e)", "iso27001": "A.8.8", "soc2": "CC7.1", "nist_csf": "PR.IP-12",
      "cis": "CIS 7.3", "gdpr": "Art.32(1)(a)", "dora": "Art.9(3)(a)", "cyber_essentials": "Security update management"}),

    # ── Incident notification ──
    ("NOT-01", "Incident Notification", "Regulatory notification process",
     "A process exists to notify regulators within required timeframes",
     6, "Is there a process to notify regulators (e.g. within 72h for GDPR, 24h for NIS2)?",
     {"nis2": "Art.23(1)", "iso27001": "A.5.26", "soc2": "CC7.5", "nist_csf": "RS.CO-02",
      "cis": "CIS 17.1", "gdpr": "Art.33", "dora": "Art.19(1)", "cyber_essentials": "—"}),

    ("NOT-02", "Incident Notification", "Customer/data subject notification",
     "Affected individuals are notified when their data is compromised",
     4, "Is there a process to notify affected individuals after a data breach?",
     {"nis2": "Art.23(2)", "iso27001": "A.5.26", "soc2": "CC2.3", "nist_csf": "RS.CO-03",
      "cis": "—", "gdpr": "Art.34", "dora": "Art.19(3)", "cyber_essentials": "—"}),
]


# ══════════════════════════════════════════════════
# Compliance assessor
# ══════════════════════════════════════════════════

class ComplianceAssessor:
    """Runs all compliance checks and produces a structured report."""

    def __init__(self, config, modules: dict = None):
        self.cfg = config
        self.mods = modules or {}

    def assess(self, frameworks: list = None) -> dict:
        """Run all controls. frameworks=list of framework IDs to include (None=all)."""
        selected = frameworks or ALL_FRAMEWORK_IDS

        # Load declarative answers from DB
        answers = {}
        try:
            from core.database import ComplianceAnswer
            for a in ComplianceAnswer.select():
                answers[a.control_id] = {"answer": a.answer, "detail": a.detail or "",
                                         "by": a.answered_by or "", "at": a.updated_at.isoformat() if a.updated_at else ""}
        except Exception as e:
            logger.debug("Failed to load compliance answers from DB: %s", e)

        auto_results = []
        decl_results = []
        auto_weight = 0
        auto_passed = 0
        decl_weight = 0
        decl_passed = 0

        # ── Automated controls ──
        for ctrl in CONTROLS:
            cid, cat, title, desc, check_fn, weight, rec, mappings = ctrl
            auto_weight += weight
            status, detail = "FAIL", ""
            try:
                fn = getattr(self, check_fn, None)
                if fn:
                    ok, detail = fn()
                    status = "PASS" if ok else "FAIL"
            except Exception as e:
                detail = f"Error: {e}"
            if status == "PASS":
                auto_passed += weight
            filtered = {k: v for k, v in mappings.items() if k in selected}
            auto_results.append({
                "id": cid, "category": cat, "title": title,
                "description": desc, "status": status, "detail": detail,
                "weight": weight, "recommendation": rec,
                "mappings": filtered, "type": "automated",
            })

        # ── Declarative controls ──
        for ctrl in DECLARATIVE_CONTROLS:
            cid, cat, title, desc, weight, prompt, mappings = ctrl
            decl_weight += weight
            ans = answers.get(cid, {})
            answer_val = ans.get("answer", "unanswered")
            ans_detail = ans.get("detail", "")
            ans_by = ans.get("by", "")

            if answer_val == "yes":
                status = "PASS"
            elif answer_val == "partial":
                status = "PARTIAL"
                decl_passed += weight // 2
            elif answer_val == "no":
                status = "FAIL"
            else:
                status = "UNANSWERED"

            if answer_val == "yes":
                decl_passed += weight

            detail_parts = []
            if ans_detail:
                detail_parts.append(ans_detail)
            if ans_by:
                detail_parts.append(f"Answered by: {ans_by}")

            filtered = {k: v for k, v in mappings.items() if k in selected}
            decl_results.append({
                "id": cid, "category": cat, "title": title,
                "description": desc, "status": status,
                "detail": " | ".join(detail_parts) if detail_parts else "",
                "weight": weight, "recommendation": prompt,
                "mappings": filtered, "type": "declarative",
                "prompt": prompt, "answer": answer_val,
                "answer_detail": ans_detail, "answered_by": ans_by,
            })

        # ── Combined scoring ──
        all_results = auto_results + decl_results
        total_weight = auto_weight + decl_weight
        total_passed = auto_passed + decl_passed

        auto_score = round(auto_passed / max(auto_weight, 1) * 100)
        decl_answered = sum(1 for r in decl_results if r["status"] != "UNANSWERED")
        decl_score = round(decl_passed / max(decl_weight, 1) * 100) if decl_answered > 0 else None
        combined_score = round(total_passed / max(total_weight, 1) * 100)

        categories = {}
        for r in all_results:
            cat = r["category"]
            if cat not in categories:
                categories[cat] = {"pass": 0, "fail": 0, "unanswered": 0, "total": 0, "controls": []}
            categories[cat]["total"] += 1
            categories[cat]["controls"].append(r)
            if r["status"] == "PASS":
                categories[cat]["pass"] += 1
            elif r["status"] == "UNANSWERED":
                categories[cat]["unanswered"] += 1
            else:
                categories[cat]["fail"] += 1

        recommendations = sorted(
            [r for r in all_results if r["status"] not in ("PASS",)],
            key=lambda x: (-x["weight"], x["status"] == "UNANSWERED"))

        risk = "LOW" if combined_score >= 80 else "MEDIUM" if combined_score >= 60 else "HIGH" if combined_score >= 40 else "CRITICAL"
        selected_fw = {k: FRAMEWORKS[k] for k in selected if k in FRAMEWORKS}

        return {
            "score": combined_score, "risk_level": risk,
            "auto_score": auto_score, "declarative_score": decl_score,
            "generated_at": datetime.now().isoformat(),
            "total_controls": len(all_results),
            "auto_controls": len(auto_results),
            "declarative_controls": len(decl_results),
            "declarative_answered": decl_answered,
            "passed": sum(1 for r in all_results if r["status"] == "PASS"),
            "failed": sum(1 for r in all_results if r["status"] in ("FAIL", "PARTIAL")),
            "unanswered": sum(1 for r in all_results if r["status"] == "UNANSWERED"),
            "categories": categories,
            "controls": all_results,
            "recommendations": recommendations,
            "frameworks": selected_fw,
            "available_frameworks": FRAMEWORKS,
        }

    # ── Individual checks (same as before) ──

    def check_sniffer_active(self):
        s = self.mods.get("sniffer")
        if not s: return False, "Sniffer module not loaded"
        st = s.stats
        if st.get("running"): return True, f"Active: {st.get('pps',0)} pkt/s, {st.get('packets',0)} total"
        return False, "Sniffer thread is not running"

    def check_host_inventory(self):
        from core.database import Host
        c = Host.select().count()
        return (True, f"{c} hosts") if c >= 1 else (False, "No hosts discovered")

    def check_dns_monitoring(self):
        from core.database import DnsLog
        c = DnsLog.select().where(DnsLog.ts >= datetime.now() - timedelta(hours=24)).count()
        return (True, f"{c} DNS queries in 24h") if c > 0 else (False, "No DNS queries logged")

    def check_threat_engine(self):
        e = self.mods.get("engine")
        return (True, "Active") if e else (False, "Not loaded")

    def check_advanced_detectors(self):
        o = self.mods.get("orchestrator")
        if not o: return False, "Orchestrator not loaded"
        active = sum(1 for d in o.detectors if d.status == "active")
        observing = sum(1 for d in o.detectors if d.status == "observing")
        return (True, f"{active} active, {observing} observing") if active >= 4 else (False, f"Only {active} active")

    def check_threat_feeds(self):
        tf = self.mods.get("threat_feeds")
        if not tf: return False, "Not loaded"
        st = tf.stats
        return (True, f"{st['ips']} IPs, {st['domains']} domains") if st.get("ips", 0) > 0 else (False, "No indicators loaded")

    def check_killchain(self):
        return (True, "Active") if self.mods.get("killchain") else (False, "Not loaded")

    def check_firewall_active(self):
        d = self.mods.get("defense")
        if not d: return False, "Not loaded"
        return (True, f"Backend: {d._fw_backend}") if d._fw_backend != "none" else (False, "No firewall detected")

    def check_auto_block(self):
        d = self.mods.get("defense")
        if not d: return False, "Not loaded"
        return (True, "Enabled") if d.auto_block else (False, "Disabled")

    def check_escalation(self):
        d = self.mods.get("defense")
        return (True, "Active") if d and hasattr(d, '_escalation') else (False, "Not available")

    def check_incident_workflow(self):
        if self.cfg.get("email.enabled"):
            return True, f"Enabled (timeout: {self.cfg.get('email.approval_timeout_minutes',0)}min)"
        return False, "Email disabled"

    def check_forensic_enabled(self):
        return (True, "Enabled") if self.cfg.get("client_agent.collect_after_incident") else (False, "Disabled")

    def check_admin_emails(self):
        emails = self.cfg.get("email.admin_emails", [])
        return (True, f"{len(emails)} admin(s)") if emails else (False, "No admin email configured")

    def check_auth_configured(self):
        from core.database import WebUser
        c = WebUser.select().count()
        return (True, f"{c} account(s)") if c > 0 else (False, "No accounts")

    def check_2fa(self):
        from core.database import WebUser
        c = WebUser.select().where(WebUser.totp_secret.is_null(False), WebUser.totp_secret != "", WebUser.role == "admin").count()
        return (True, f"{c} admin(s) with 2FA") if c > 0 else (False, "No admin has 2FA")

    def check_login_lockout(self):
        m = self.cfg.get("web.max_login_attempts", 0)
        return (True, f"Lockout after {m} attempts") if m and m <= 10 else (False, "Not configured")

    def check_tls(self):
        c = self.cfg.get("web.ssl_cert", "")
        return (True, f"Certificate: {c}") if c and os.path.exists(c) else (False, "No TLS certificate")

    def check_ssh(self):
        try:
            from core.hardening import SSHHardener
            r = SSHHardener().verify()
            return (True, "Hardened") if r.get("secure") else (False, "; ".join(r.get("issues", [])))
        except Exception as e:
            return False, str(e)

    def check_integrity(self):
        try:
            from core.hardening import IntegrityCheck
            r = IntegrityCheck.verify()
            return (True, f"{r.get('total_checked',0)} files OK") if r.get("ok") else (False, f"{len(r.get('modified',[]))} modified")
        except Exception as e:
            return False, str(e)

    def check_backup(self):
        d = self.cfg.get("backup.directory", "")
        if d and os.path.isdir(d):
            f = [x for x in os.listdir(d) if x.endswith(".tar.gz")]
            return (True, f"{len(f)} backup(s)") if f else (False, "No backups found")
        return False, "Backup directory not configured"

    def check_audit_chain(self):
        ac = self.mods.get("audit_chain")
        if not ac: return False, "Not loaded"
        try:
            v = ac.verify()
            return (True, f"Intact ({v.get('entries',0)} entries)") if v.get("ok") else (False, "Chain integrity FAILED")
        except Exception as e:
            return False, str(e)

    def check_retention(self):
        a = self.cfg.get("retention.alerts_days", 0)
        return (True, f"Alerts: {a}d") if a > 0 else (False, "Not configured")

    def check_honeypot(self):
        hp = self.mods.get("honeypot")
        return (True, f"{len(hp.ports)} ports") if hp and hp.enabled else (False, "Disabled")


# ══════════════════════════════════════════════════
# PDF generator
# ══════════════════════════════════════════════════

def _header_footer(canvas, doc, company, report_id):
    """Professional header and footer on every page (except cover)."""
    if doc.page == 1:
        return  # No header/footer on cover
    canvas.saveState()
    w, h = A4

    # Header: blue line + company + report ID
    canvas.setStrokeColor(C_BLUE)
    canvas.setLineWidth(1.5)
    canvas.line(18*mm, h - 14*mm, w - 18*mm, h - 14*mm)

    canvas.setFont("Helvetica-Bold", 7)
    canvas.setFillColor(C_DARK)
    canvas.drawString(18*mm, h - 12*mm, company.upper() if company else "CGS COMPLIANCE REPORT")

    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(C_GRAY)
    canvas.drawRightString(w - 18*mm, h - 12*mm, report_id)

    # Footer: line + page number + date
    canvas.setStrokeColor(C_BORDER)
    canvas.setLineWidth(0.5)
    canvas.line(18*mm, 14*mm, w - 18*mm, 14*mm)

    canvas.setFont("Helvetica", 6.5)
    canvas.setFillColor(C_GRAY)
    canvas.drawString(18*mm, 10*mm, "CGS — Autonomous Micro-SIEM")
    canvas.drawCentredString(w / 2, 10*mm, f"Page {doc.page}")
    canvas.drawRightString(w - 18*mm, 10*mm, "CONFIDENTIAL")

    canvas.restoreState()


def generate_compliance_pdf(assessment: dict, config=None,
                            output_path: str = None, company: str = "") -> bytes:
    styles = _styles()
    buf = io.BytesIO()
    now = datetime.now()
    report_id = f"CGS-CR-{now.strftime('%Y%m%d-%H%M')}"

    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=20*mm, bottomMargin=20*mm,
                            title="CGS Compliance Report", author="CGS")

    story = []
    score = assessment["score"]
    risk = assessment["risk_level"]
    selected_fw = assessment.get("frameworks", FRAMEWORKS)
    fw_names = ", ".join(f["name"] for f in selected_fw.values())

    # ════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════
    # Top blue accent bar
    story.append(HRFlowable(width="100%", thickness=3, color=C_BLUE, spaceAfter=25*mm))
    story.append(Spacer(1, 10*mm))

    if company:
        story.append(Paragraph(company.upper(),
            ParagraphStyle("CO", parent=styles["cover_title"], fontSize=20,
                           textColor=C_GRAY, spaceAfter=6*mm, leading=24)))

    story.append(Paragraph("CYBERSECURITY", styles["cover_title"]))
    story.append(Paragraph("COMPLIANCE REPORT", styles["cover_title"]))
    story.append(Spacer(1, 6*mm))
    story.append(HRFlowable(width="30%", thickness=2, color=C_BLUE,
                            spaceAfter=6*mm, hAlign="CENTER"))
    story.append(Paragraph("Security Posture Assessment", styles["cover_sub"]))
    story.append(Spacer(1, 2*mm))
    story.append(Paragraph(fw_names,
        ParagraphStyle("FW", parent=styles["cover_sub"], fontSize=8,
                       textColor=C_LIGHT_BLUE, leading=12)))
    story.append(Spacer(1, 12*mm))

    # Determine if all declarative controls are answered and all controls pass
    all_answered = assessment.get("unanswered", 0) == 0
    all_pass = score == 100
    fully_compliant = all_answered and all_pass

    score_color = C_GREEN if score >= 80 else C_YELLOW if score >= 60 else C_ORANGE if score >= 40 else C_RED
    rc = {"LOW": C_GREEN, "MEDIUM": C_YELLOW, "HIGH": C_ORANGE, "CRITICAL": C_RED}.get(risk, C_GRAY)
    grade = "EXCELLENT" if score >= 90 else "GOOD" if score >= 80 else "SATISFACTORY" if score >= 60 else "INSUFFICIENT" if score >= 40 else "CRITICAL"

    # Score box: visual centered box with score + grade + risk
    score_box_data = [[
        Paragraph(f"<font color='{score_color.hexval()}' size='36'><b>{score}</b></font>"
                  f"<font color='{C_GRAY.hexval()}' size='14'>/100</font>",
                  ParagraphStyle("SB", alignment=TA_CENTER)),
        Paragraph(f"<font size='10' color='{C_GRAY.hexval()}'>COMPLIANCE GRADE</font><br/>"
                  f"<font color='{score_color.hexval()}' size='20'><b>{grade}</b></font><br/>"
                  f"<font size='9' color='{C_GRAY.hexval()}'>Risk Level: </font>"
                  f"<font size='9' color='{rc.hexval()}'><b>{risk}</b></font>",
                  ParagraphStyle("GR", alignment=TA_CENTER, leading=16)),
    ]]
    score_t = Table(score_box_data, colWidths=[55*mm, 85*mm], hAlign="CENTER")
    score_t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BOX", (0, 0), (-1, -1), 1, C_BORDER),
        ("LINEAFTER", (0, 0), (0, -1), 0.5, C_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, -1), C_LIGHT_GRAY),
    ]))
    story.append(score_t)
    story.append(Spacer(1, 10*mm))

    # Summary table
    summary_data = []
    if company:
        summary_data.append(["Organization", company])
    summary_data += [
        ["Report Date", now.strftime("%B %d, %Y at %H:%M")],
        ["Report ID", report_id],
        ["Controls Assessed", f"{assessment['total_controls']} ({assessment.get('auto_controls',0)} automated + {assessment.get('declarative_controls',0)} self-assessed)"],
        ["Result", f"{assessment['passed']} passed, {assessment['failed']} failed" +
                   (f", {assessment.get('unanswered',0)} unanswered" if assessment.get("unanswered") else "")],
        ["Frameworks", fw_names],
    ]
    st = Table(summary_data, colWidths=[42*mm, 100*mm], hAlign="CENTER")
    st.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), C_BLUE), ("TEXTCOLOR", (1, 0), (1, -1), C_DARK),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, C_BORDER),
    ]))
    story.append(st)
    story.append(Spacer(1, 15*mm))

    # Confidentiality / certification notice
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BLUE, spaceAfter=4*mm))
    if fully_compliant:
        conf = (f"CERTIFIED — This document certifies that "
                f"{company + ' ' if company else 'the assessed organization '}"
                f"achieves <b>full compliance</b> across all assessed controls. "
                f"This report may be presented to auditors, partners, regulators, or clients "
                f"as evidence of cybersecurity compliance.")
    else:
        conf = (f"CONFIDENTIAL — This document assesses the cybersecurity posture "
                f"{'of ' + company + ' ' if company else ''}"
                f"and may be shared with auditors, partners, regulators, or clients "
                f"to demonstrate compliance with applicable cybersecurity frameworks.")
    story.append(Paragraph(conf, ParagraphStyle("CONF", parent=styles["small"],
                 alignment=TA_CENTER, fontSize=7)))
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph("CGS — Autonomous Micro-SIEM", styles["cover_sub"]))
    story.append(PageBreak())

    # ════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════
    story.append(Paragraph("1. Executive Summary", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))

    passed = assessment["passed"]
    total = assessment["total_controls"]
    failed = assessment["failed"]
    recs = assessment["recommendations"]
    org = f"<b>{company}</b>" if company else "The assessed organization"

    auto_score = assessment.get("auto_score", 0)
    decl_score = assessment.get("declarative_score")
    n_auto = assessment.get("auto_controls", 0)
    n_decl = assessment.get("declarative_controls", 0)
    n_unans = assessment.get("unanswered", 0)

    story.append(Paragraph(
        f"{org} has been assessed against <b>{len(selected_fw)}</b> cybersecurity "
        f"frameworks. The assessment covers <b>{total}</b> security controls: "
        f"<b>{n_auto}</b> verified automatically by the CGS monitoring system and "
        f"<b>{n_decl}</b> self-assessed by the organization.",
        styles["body"]))
    story.append(Spacer(1, 4*mm))

    # KPI boxes
    kpi_data = [[
        Paragraph(f"<font size='18' color='{score_color.hexval()}'><b>{score}</b></font><br/>"
                  f"<font size='7' color='{C_GRAY.hexval()}'>COMBINED</font>",
                  ParagraphStyle("K", alignment=TA_CENTER, leading=14)),
        Paragraph(f"<font size='18' color='{C_LIGHT_BLUE.hexval()}'><b>{auto_score}</b></font><br/>"
                  f"<font size='7' color='{C_GRAY.hexval()}'>AUTOMATED</font>",
                  ParagraphStyle("K", alignment=TA_CENTER, leading=14)),
        Paragraph(f"<font size='18' color='{C_LIGHT_BLUE.hexval()}'><b>{decl_score if decl_score is not None else 'N/A'}</b></font><br/>"
                  f"<font size='7' color='{C_GRAY.hexval()}'>SELF-ASSESSED</font>",
                  ParagraphStyle("K", alignment=TA_CENTER, leading=14)),
        Paragraph(f"<font size='18' color='{C_GREEN.hexval()}'><b>{passed}</b></font><br/>"
                  f"<font size='7' color='{C_GRAY.hexval()}'>PASSED</font>",
                  ParagraphStyle("K", alignment=TA_CENTER, leading=14)),
        Paragraph(f"<font size='18' color='{C_RED.hexval()}'><b>{failed}</b></font><br/>"
                  f"<font size='7' color='{C_GRAY.hexval()}'>FAILED</font>",
                  ParagraphStyle("K", alignment=TA_CENTER, leading=14)),
    ]]
    kpi_t = Table(kpi_data, colWidths=[34*mm]*5)
    kpi_t.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.5, C_BORDER),
        ("LINEAFTER", (0, 0), (-2, -1), 0.3, C_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, -1), C_LIGHT_GRAY),
    ]))
    story.append(kpi_t)
    story.append(Spacer(1, 6*mm))

    if recs:
        story.append(Paragraph(
            f"<b>Top priority recommendation:</b> {recs[0]['id']} — {recs[0]['title']}",
            ParagraphStyle("TP", parent=styles["body"], textColor=C_ORANGE,
                           borderWidth=0.5, borderColor=C_ORANGE, borderPadding=6,
                           backColor=HexColor("#FFFBEB"))))
    story.append(PageBreak())

    # ════════════════════════════════════════
    # SCORE BY CATEGORY
    # ════════════════════════════════════════
    story.append(Paragraph("2. Score by Category", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))
    cat_rows = [["Category", "Passed", "Failed", "Score", "Status"]]
    cat_status_colors = []
    for i, (cat_name, cat_data) in enumerate(assessment["categories"].items()):
        cs = round(cat_data["pass"] / max(cat_data["total"], 1) * 100)
        st_text = "COMPLIANT" if cs >= 80 else "PARTIAL" if cs >= 50 else "NON-COMPLIANT"
        cat_rows.append([cat_name, str(cat_data["pass"]), str(cat_data["fail"]), f"{cs}%", st_text])
        cat_status_colors.append(C_GREEN if cs >= 80 else C_ORANGE if cs >= 50 else C_RED)
    ct = Table(cat_rows, colWidths=[50*mm, 18*mm, 18*mm, 18*mm, 42*mm])
    ts = [
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), C_BLUE), ("TEXTCOLOR", (0, 0), (-1, 0), C_WHITE),
        ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 1), (-1, -1), 0.3, C_BORDER), ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_WHITE, C_LIGHT_GRAY]),
    ]
    # Color the status column per row
    for i, color in enumerate(cat_status_colors):
        ts.append(("TEXTCOLOR", (4, i+1), (4, i+1), color))
        ts.append(("FONTNAME", (4, i+1), (4, i+1), "Helvetica-Bold"))
    ct.setStyle(TableStyle(ts))
    story.append(ct)
    story.append(PageBreak())

    # ════════════════════════════════════════
    # DETAILED CONTROLS
    # ════════════════════════════════════════
    story.append(Paragraph("3. Automated Controls (verified by CGS)", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))
    story.append(Paragraph(
        "The following controls are verified automatically and continuously by CGS. "
        "Their status reflects the actual state of the system at the time of assessment.",
        styles["body"]))

    # Auto controls grouped by category
    auto_cats = {}
    decl_cats = {}
    for cat_name, cat_data in assessment["categories"].items():
        auto_ctrls = [c for c in cat_data["controls"] if c.get("type") == "automated"]
        decl_ctrls = [c for c in cat_data["controls"] if c.get("type") == "declarative"]
        if auto_ctrls:
            auto_cats[cat_name] = auto_ctrls
        if decl_ctrls:
            decl_cats[cat_name] = decl_ctrls

    for cat_name, ctrls in auto_cats.items():
        story.append(Paragraph(cat_name, styles["h2"]))
        for ctrl in ctrls:
            bg = C_BG_GREEN if ctrl["status"] == "PASS" else C_BG_RED
            st_style = styles["pass"] if ctrl["status"] == "PASS" else styles["fail"]
            rows = [
                [Paragraph(f"<b>{ctrl['id']}</b>", styles["label"]),
                 Paragraph(f"<b>{ctrl['title']}</b>", styles["value"]),
                 Paragraph(ctrl["status"], st_style)],
                ["", Paragraph(ctrl["description"], styles["small"]), ""],
            ]
            if ctrl["detail"]:
                rows.append(["", Paragraph(f"<i>{ctrl['detail']}</i>", styles["small"]), ""])
            # Framework references
            fw_refs = " | ".join(f"{FRAMEWORKS[k]['name']}: {v}" for k, v in ctrl["mappings"].items() if v != "—")
            if fw_refs:
                rows.append(["", Paragraph(fw_refs, ParagraphStyle("FR", parent=styles["small"], textColor=C_LIGHT_BLUE, fontSize=7)), ""])
            if ctrl["status"] != "PASS":
                rows.append(["", Paragraph(f"<b>Recommendation:</b> {ctrl['recommendation']}", styles["rec"]), ""])
            t = Table(rows, colWidths=[18*mm, 125*mm, 18*mm])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("BACKGROUND", (0, 0), (-1, 0), bg),
                ("LINEBELOW", (0, -1), (-1, -1), 0.5, C_BORDER),
            ]))
            story.append(t)
            story.append(Spacer(1, 2*mm))
    story.append(PageBreak())

    # ════════════════════════════════════════
    # DECLARATIVE CONTROLS (self-assessed)
    # ════════════════════════════════════════
    if fully_compliant:
        story.append(Paragraph("4. Organizational Controls (declared by the organization)", styles["h1"]))
    else:
        story.append(Paragraph("4. Self-Assessed Controls (declared by the organization)", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))

    if not all_answered:
        story.append(Paragraph(
            "<b>DISCLAIMER:</b> The following controls are self-assessed by the organization. "
            "Their status is based on declarations provided by authorized personnel and "
            "<b>has not been independently verified</b>. An external audit may be required "
            "to confirm these assessments. "
            f"<b>{assessment.get('unanswered', 0)} control(s) have not yet been answered.</b>",
            ParagraphStyle("DISCL", parent=styles["body"], fontSize=9,
                           textColor=C_ORANGE, backColor=HexColor("#FFFBEB"),
                           borderWidth=0.5, borderColor=C_ORANGE, borderPadding=6)))
        story.append(Spacer(1, 4*mm))
    elif not all_pass:
        story.append(Paragraph(
            "<b>NOTE:</b> The following controls have been assessed by the organization. "
            "All questions have been answered. Some controls require remediation "
            "to achieve full compliance.",
            ParagraphStyle("NOTE", parent=styles["body"], fontSize=9,
                           textColor=C_YELLOW, backColor=HexColor("#FFFBEB"),
                           borderWidth=0.5, borderColor=C_YELLOW, borderPadding=6)))
        story.append(Spacer(1, 4*mm))
    # If fully_compliant: no disclaimer at all

    for cat_name, ctrls in decl_cats.items():
        story.append(Paragraph(cat_name, styles["h2"]))
        for ctrl in ctrls:
            ans = ctrl.get("answer", "unanswered")
            if ctrl["status"] == "PASS":
                bg, st_style = C_BG_GREEN, styles["pass"]
            elif ctrl["status"] == "UNANSWERED":
                bg = C_LIGHT_GRAY
                st_style = ParagraphStyle("UA", parent=styles["small"],
                                          textColor=C_GRAY, fontName="Helvetica-Bold")
            else:
                bg, st_style = C_BG_RED, styles["fail"]

            rows = [
                [Paragraph(f"<b>{ctrl['id']}</b>", styles["label"]),
                 Paragraph(f"<b>{ctrl['title']}</b>", styles["value"]),
                 Paragraph(ctrl["status"], st_style)],
                ["", Paragraph(ctrl.get("prompt", ctrl["description"]), styles["small"]), ""],
            ]
            if ctrl.get("answer_detail"):
                rows.append(["", Paragraph(f"Evidence: <i>{ctrl['answer_detail']}</i>", styles["small"]), ""])
            if ctrl.get("answered_by"):
                rows.append(["", Paragraph(f"Declared by: {ctrl['answered_by']}", styles["small"]), ""])
            fw_refs = " | ".join(f"{FRAMEWORKS[k]['name']}: {v}" for k, v in ctrl["mappings"].items() if v != "—")
            if fw_refs:
                rows.append(["", Paragraph(fw_refs, ParagraphStyle("FR2", parent=styles["small"],
                            textColor=C_LIGHT_BLUE, fontSize=7)), ""])

            t = Table(rows, colWidths=[18*mm, 125*mm, 18*mm])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("BACKGROUND", (0, 0), (-1, 0), bg),
                ("LINEBELOW", (0, -1), (-1, -1), 0.5, C_BORDER),
            ]))
            story.append(t)
            story.append(Spacer(1, 2*mm))

    if not decl_cats:
        story.append(Paragraph("No self-assessed controls configured.", styles["body"]))
    story.append(PageBreak())

    # ════════════════════════════════════════
    # FRAMEWORK ALIGNMENT PAGES
    # ════════════════════════════════════════
    section_num = 5
    for fw_id, fw_info in selected_fw.items():
        story.append(Paragraph(f"{section_num}. {fw_info['name']} Alignment", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))
        story.append(Paragraph(f"<b>{fw_info['full']}</b> ({fw_info['version']})", styles["body"]))

        fw_rows = [[fw_info["name"], "Control", "Status"]]
        for ctrl in assessment["controls"]:
            ref = ctrl["mappings"].get(fw_id, "")
            if ref and ref != "—":
                fw_rows.append([ref, f"{ctrl['id']} — {ctrl['title']}", ctrl["status"]])

        if len(fw_rows) > 1:
            ft = Table(fw_rows, colWidths=[35*mm, 100*mm, 25*mm])
            ft.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), C_BLUE), ("TEXTCOLOR", (0, 0), (-1, 0), C_WHITE),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LINEBELOW", (0, 1), (-1, -1), 0.3, C_BORDER), ("ALIGN", (2, 0), (2, -1), "CENTER"),
            ]))
            story.append(ft)
        else:
            story.append(Paragraph("No controls mapped to this framework.", styles["body"]))

        section_num += 1
        story.append(PageBreak())

    # ════════════════════════════════════════
    # RECOMMENDATIONS
    # ════════════════════════════════════════
    story.append(Paragraph(f"{section_num}. Prioritized Recommendations", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))
    for rec in recs:
        pr = "HIGH" if rec["weight"] >= 8 else "MEDIUM" if rec["weight"] >= 6 else "LOW"
        pc = C_RED if pr == "HIGH" else C_ORANGE if pr == "MEDIUM" else C_YELLOW
        story.append(Paragraph(
            f"<font color='{pc.hexval()}'><b>{pr}</b></font> — <b>{rec['id']}: {rec['title']}</b>",
            styles["h3"]))
        story.append(Paragraph(rec["recommendation"], styles["rec"]))
        fw_refs = ", ".join(f"{FRAMEWORKS[k]['name']} {v}" for k, v in rec["mappings"].items() if v != "—")
        if fw_refs:
            story.append(Paragraph(f"<i>{fw_refs}</i>",
                ParagraphStyle("fw", parent=styles["small"], leftIndent=10*mm)))
        story.append(Spacer(1, 3*mm))
    if not recs:
        story.append(Paragraph("<font color='#16A34A'><b>All controls passed.</b></font>", styles["body"]))
    section_num += 1
    story.append(PageBreak())

    # ════════════════════════════════════════
    # NETWORK METRICS
    # ════════════════════════════════════════
    story.append(Paragraph(f"{section_num}. Network Security Metrics", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))
    try:
        from core.database import Alert, Host
        today = datetime.now().replace(hour=0, minute=0, second=0)
        week_ago = datetime.now() - timedelta(days=7)
        md = [["Metric", "Value"],
              ["Total hosts monitored", str(Host.select().count())],
              ["Active hosts", str(Host.select().where(Host.status == "up").count())],
              ["Hosts with risk > 50", str(Host.select().where(Host.risk_score > 50).count())],
              ["Alerts today", str(Alert.select().where(Alert.ts >= today).count())],
              ["Critical alerts (7d)", str(Alert.select().where(Alert.severity <= 2, Alert.ts >= week_ago).count())],
              ["Unacknowledged alerts", str(Alert.select().where(Alert.ack == False).count())]]
        mt = Table(md, colWidths=[80*mm, 60*mm])
        mt.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BACKGROUND", (0, 0), (-1, 0), C_LIGHT_GRAY),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, C_BORDER), ("ALIGN", (1, 1), (1, -1), "CENTER"),
        ]))
        story.append(mt)
    except Exception as e:
        story.append(Paragraph(f"Metrics unavailable: {e}", styles["body"]))
    section_num += 1

    # ════════════════════════════════════════
    # GRC SUMMARY
    # ════════════════════════════════════════
    grc = assessment.get("grc")
    if grc and not grc.get("error"):
        story.append(PageBreak())
        story.append(Paragraph(f"{section_num}. Governance, Risk &amp; Compliance Overview", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4*mm))

        # Risk summary
        rk = grc.get("risks", {})
        if rk.get("total", 0) > 0:
            story.append(Paragraph("Risk Register", styles["h2"]))
            rk_data = [["Metric", "Value"],
                       ["Total risks identified", str(rk["total"])],
                       ["Open risks", str(rk.get("open", 0))],
                       ["Critical risks (score >= 15)", str(rk.get("critical", 0))]]
            rkt = Table(rk_data, colWidths=[80*mm, 50*mm])
            rkt.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), C_LIGHT_GRAY),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LINEBELOW", (0, 0), (-1, -1), 0.3, C_BORDER), ("ALIGN", (1, 1), (1, -1), "CENTER"),
            ]))
            story.append(rkt)
            story.append(Spacer(1, 3*mm))

            top5 = rk.get("top5", [])
            if top5:
                story.append(Paragraph("Top risks by score:", styles["body"]))
                t5_data = [["Risk", "Score", "Status"]]
                for tr in top5:
                    t5_data.append([tr["title"], str(tr["score"]), tr["status"]])
                t5t = Table(t5_data, colWidths=[95*mm, 25*mm, 30*mm])
                t5t.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("BACKGROUND", (0, 0), (-1, 0), C_BLUE), ("TEXTCOLOR", (0, 0), (-1, 0), C_WHITE),
                    ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LINEBELOW", (0, 1), (-1, -1), 0.3, C_BORDER), ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ]))
                story.append(t5t)
            story.append(Spacer(1, 4*mm))

        # Asset summary
        ast = grc.get("assets", {})
        if ast.get("total", 0) > 0:
            story.append(Paragraph("Asset Inventory", styles["h2"]))
            by_type = ast.get("by_type", {})
            ast_rows = [["Asset Type", "Count"]]
            for atype, cnt in sorted(by_type.items()):
                ast_rows.append([atype.capitalize(), str(cnt)])
            ast_rows.append(["Total", str(ast["total"])])
            astt = Table(ast_rows, colWidths=[80*mm, 40*mm])
            astt.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), C_LIGHT_GRAY),
                ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LINEBELOW", (0, 0), (-1, -1), 0.3, C_BORDER), ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ]))
            story.append(astt)
            story.append(Spacer(1, 4*mm))

        # Policy summary
        pol = grc.get("policies", {})
        if pol.get("total", 0) > 0:
            story.append(Paragraph("Policy Status", styles["h2"]))
            story.append(Paragraph(
                f"<b>{pol['total']}</b> policies registered: "
                f"<b>{pol.get('approved', 0)}</b> approved, "
                f"<b>{pol.get('draft', 0)}</b> in draft.",
                styles["body"]))

        # Vendor summary
        vnd = grc.get("vendors", {})
        if vnd.get("total", 0) > 0:
            story.append(Paragraph("Third-Party Risk", styles["h2"]))
            story.append(Paragraph(
                f"<b>{vnd['total']}</b> vendors assessed: "
                f"<b>{vnd.get('high_risk', 0)}</b> rated high/critical risk.",
                styles["body"]))

        # Audit findings
        af = grc.get("audit_findings_open", 0)
        if af > 0:
            story.append(Paragraph("Open Audit Findings", styles["h2"]))
            story.append(Paragraph(
                f"<font color='{C_ORANGE.hexval()}'><b>{af}</b></font> audit finding(s) "
                f"remain open and require remediation.",
                styles["body"]))

        # Compliance trend
        trend = grc.get("compliance_trend", {})
        if any(v is not None for v in trend.values()):
            story.append(Paragraph("Compliance Trend", styles["h2"]))
            trend_rows = [["Period", "Score Delta"]]
            labels = {"m1": "vs. 1 month ago", "m3": "vs. 3 months ago",
                      "m6": "vs. 6 months ago", "m12": "vs. 12 months ago"}
            for k, label in labels.items():
                v = trend.get(k)
                if v is not None:
                    prefix = "+" if v > 0 else ""
                    trend_rows.append([label, f"{prefix}{v} points"])
            if len(trend_rows) > 1:
                trt = Table(trend_rows, colWidths=[80*mm, 50*mm])
                trt.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), C_LIGHT_GRAY),
                    ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LINEBELOW", (0, 0), (-1, -1), 0.3, C_BORDER), ("ALIGN", (1, 1), (1, -1), "CENTER"),
                ]))
                story.append(trt)

        section_num += 1

    # ════════════════════════════════════════
    # ATTESTATION
    # ════════════════════════════════════════
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=6*mm))
    story.append(Paragraph("COMPLIANCE ATTESTATION", styles["h1"]))

    org_name = f"<b>{company}</b>" if company else "the assessed organization"
    grade = "EXCELLENT" if score >= 90 else "GOOD" if score >= 80 else "SATISFACTORY" if score >= 60 else "INSUFFICIENT" if score >= 40 else "CRITICAL"
    gc = C_GREEN if score >= 80 else C_YELLOW if score >= 60 else C_ORANGE if score >= 40 else C_RED

    story.append(Paragraph(
        f"This report certifies that {org_name} has undergone an automated cybersecurity "
        f"compliance assessment on <b>{now.strftime('%B %d, %Y')}</b>, covering "
        f"<b>{total}</b> security controls aligned with <b>{fw_names}</b>."
        f"<br/><br/>"
        f"The overall compliance score is "
        f"<font color='{score_color.hexval()}'><b>{score}/100</b></font> — "
        f"Grade: <font color='{gc.hexval()}'><b>{grade}</b></font>."
        f"<br/><br/>"
        f"<b>{passed}</b> out of <b>{total}</b> controls passed.",
        styles["body"]))

    sig_data = []
    if company:
        sig_data.append(["Organization:", company])
    sig_data += [
        ["Assessment Date:", now.strftime("%Y-%m-%d %H:%M:%S")],
        ["Report ID:", f"CGS-CR-{now.strftime('%Y%m%d-%H%M')}"],
        ["Frameworks:", fw_names],
        ["Method:", "Automated continuous monitoring & control verification"],
    ]
    sig_t = Table(sig_data, colWidths=[45*mm, 110*mm])
    sig_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), C_GRAY),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, C_BORDER),
    ]))
    story.append(Spacer(1, 4*mm))
    story.append(sig_t)
    story.append(Spacer(1, 10*mm))
    if fully_compliant:
        story.append(Paragraph(
            "All automated and organizational controls have been assessed and are in compliance. "
            "This document serves as a certificate of cybersecurity compliance and may be "
            "presented to any third party. Regular reassessment is recommended to maintain compliance.",
            ParagraphStyle("DISC", parent=styles["small"], alignment=TA_CENTER)))
    else:
        story.append(Paragraph(
            "This assessment reflects the state of security controls at the time of generation. "
            "Self-assessed controls are declared by the organization and have not been independently verified. "
            "Regular reassessment is recommended.",
            ParagraphStyle("DISC", parent=styles["small"], alignment=TA_CENTER)))

    story.append(Spacer(1, 8*mm))
    story.append(HRFlowable(width="100%", thickness=0.3, color=C_BORDER, spaceAfter=3*mm))
    story.append(Paragraph(
        f"CGS Compliance Report — {now.strftime('%Y-%m-%d %H:%M:%S')} — "
        f"Auto-generated — Review by qualified personnel recommended.",
        styles["footer"]))

    def _on_page(canvas, doc_obj):
        _header_footer(canvas, doc_obj, company, report_id)

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    pdf_bytes = buf.getvalue()
    buf.close()

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(pdf_bytes)

    return pdf_bytes


# ══════════════════════════════════════════════════
# DOCX generator (editable format)
# ══════════════════════════════════════════════════

def generate_compliance_docx(assessment: dict, config=None, company: str = "") -> bytes:
    """Generate an editable DOCX compliance report. Returns bytes."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    now = datetime.now()
    score = assessment["score"]
    risk = assessment["risk_level"]
    selected_fw = assessment.get("frameworks", FRAMEWORKS)
    fw_names = ", ".join(f["name"] for f in selected_fw.values())
    report_id = f"CGS-CR-{now.strftime('%Y%m%d-%H%M')}"
    passed = assessment["passed"]
    total = assessment["total_controls"]
    failed = assessment["failed"]
    n_unans = assessment.get("unanswered", 0)
    all_answered = n_unans == 0
    fully_compliant = all_answered and score == 100
    grade = "EXCELLENT" if score >= 90 else "GOOD" if score >= 80 else "SATISFACTORY" if score >= 60 else "INSUFFICIENT" if score >= 40 else "CRITICAL"

    # Colors
    BLUE = RGBColor(0x1E, 0x40, 0xAF)
    GREEN = RGBColor(0x16, 0xA3, 0x4A)
    RED = RGBColor(0xDC, 0x26, 0x26)
    ORANGE = RGBColor(0xEA, 0x58, 0x0C)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    DARK = RGBColor(0x0F, 0x17, 0x2A)
    LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        # Header
        header = s.header
        hp = header.paragraphs[0]
        hp.text = f"{company.upper() if company else 'CGS COMPLIANCE REPORT'}    |    {report_id}"
        hp.style.font.size = Pt(7)
        hp.style.font.color.rgb = GRAY
        # Footer
        footer = s.footer
        fp = footer.paragraphs[0]
        fp.text = "CGS — Autonomous Micro-SIEM    |    CONFIDENTIAL"
        fp.style.font.size = Pt(7)
        fp.style.font.color.rgb = GRAY

    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = BLUE if level <= 2 else DARK
        return h

    def _para(text, bold=False, color=None, size=None, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
        if align:
            p.alignment = align
        return p

    def _table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
        return t

    # ════════════════════════════════════════
    # COVER
    # ════════════════════════════════════════
    doc.add_paragraph()  # spacer
    if company:
        _para(company.upper(), bold=True, color=GRAY, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para("CYBERSECURITY COMPLIANCE REPORT", bold=True, color=DARK, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para("Security Posture Assessment", color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(fw_names, color=BLUE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{score}/100")
    run.bold = True
    run.font.size = Pt(36)
    sc = GREEN if score >= 80 else ORANGE if score >= 60 else RED
    run.font.color.rgb = sc

    _para(f"Grade: {grade}  |  Risk Level: {risk}", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Summary table
    info = []
    if company:
        info.append(("Organization", company))
    info += [
        ("Report Date", now.strftime("%B %d, %Y at %H:%M")),
        ("Report ID", report_id),
        ("Controls", f"{total} ({assessment.get('auto_controls',0)} automated + {assessment.get('declarative_controls',0)} self-assessed)"),
        ("Result", f"{passed} passed, {failed} failed" + (f", {n_unans} unanswered" if n_unans else "")),
        ("Frameworks", fw_names),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.style = 'Light List'
    for i, (label, value) in enumerate(info):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = value
        for run in t.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for run in t.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_paragraph()
    if fully_compliant:
        _para("CERTIFIED — This document certifies full compliance across all assessed controls.",
              bold=True, color=GREEN, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _para("CONFIDENTIAL — This document may be shared with auditors, partners, or regulators.",
              color=GRAY, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════
    _heading("1. Executive Summary")
    org = company if company else "The assessed organization"
    doc.add_paragraph(
        f"{org} has been assessed against {len(selected_fw)} cybersecurity frameworks. "
        f"The assessment covers {total} security controls: "
        f"{assessment.get('auto_controls',0)} verified automatically and "
        f"{assessment.get('declarative_controls',0)} self-assessed.")

    recs = assessment["recommendations"]
    if recs:
        p = doc.add_paragraph()
        run = p.add_run(f"Top priority: {recs[0]['id']} — {recs[0]['title']}")
        run.bold = True
        run.font.color.rgb = ORANGE

    doc.add_page_break()

    # ════════════════════════════════════════
    # SCORE BY CATEGORY
    # ════════════════════════════════════════
    _heading("2. Score by Category")
    cat_rows = []
    for cat_name, cat_data in assessment["categories"].items():
        cs = round(cat_data["pass"] / max(cat_data["total"], 1) * 100)
        st = "COMPLIANT" if cs >= 80 else "PARTIAL" if cs >= 50 else "NON-COMPLIANT"
        cat_rows.append([cat_name, str(cat_data["pass"]), str(cat_data["fail"]), f"{cs}%", st])
    _table(["Category", "Passed", "Failed", "Score", "Status"], cat_rows)

    doc.add_page_break()

    # ════════════════════════════════════════
    # AUTOMATED CONTROLS
    # ════════════════════════════════════════
    _heading("3. Automated Controls (verified by CGS)")
    doc.add_paragraph(
        "The following controls are verified automatically and continuously. "
        "Their status reflects the actual system state at assessment time.")

    auto_ctrls = [c for c in assessment["controls"] if c.get("type") == "automated"]
    for ctrl in auto_ctrls:
        p = doc.add_paragraph()
        run = p.add_run(f"{ctrl['id']}  ")
        run.bold = True
        run.font.color.rgb = BLUE
        run = p.add_run(ctrl["title"])
        run.bold = True
        run = p.add_run(f"  —  {ctrl['status']}")
        run.font.color.rgb = GREEN if ctrl["status"] == "PASS" else RED
        run.bold = True
        if ctrl["detail"]:
            doc.add_paragraph(ctrl["detail"]).style.font.size = Pt(9)
        if ctrl["status"] != "PASS":
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"Recommendation: {ctrl['recommendation']}")
            run2.italic = True
            run2.font.size = Pt(9)
            run2.font.color.rgb = ORANGE

    doc.add_page_break()

    # ════════════════════════════════════════
    # DECLARATIVE CONTROLS
    # ════════════════════════════════════════
    decl_ctrls = [c for c in assessment["controls"] if c.get("type") == "declarative"]

    if fully_compliant:
        _heading("4. Organizational Controls")
    else:
        _heading("4. Self-Assessed Controls")
        if not all_answered:
            p = doc.add_paragraph()
            run = p.add_run(
                "DISCLAIMER: The following controls are self-assessed. "
                "Their status has not been independently verified. "
                f"{n_unans} control(s) have not yet been answered.")
            run.font.color.rgb = ORANGE
            run.font.size = Pt(9)
            run.bold = True

    for ctrl in decl_ctrls:
        p = doc.add_paragraph()
        run = p.add_run(f"{ctrl['id']}  ")
        run.bold = True
        run.font.color.rgb = BLUE
        run = p.add_run(ctrl["title"])
        run.bold = True
        status = ctrl["status"]
        color = GREEN if status == "PASS" else ORANGE if status == "PARTIAL" else GRAY if status == "UNANSWERED" else RED
        label = "PENDING" if status == "UNANSWERED" else status
        run = p.add_run(f"  —  {label}")
        run.font.color.rgb = color
        run.bold = True

        doc.add_paragraph(ctrl.get("prompt", ctrl["description"])).runs[0].font.size = Pt(9)
        if ctrl.get("answer_detail"):
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"Evidence: {ctrl['answer_detail']}")
            run2.italic = True
            run2.font.size = Pt(9)
        if ctrl.get("answered_by"):
            p3 = doc.add_paragraph()
            run3 = p3.add_run(f"Declared by: {ctrl['answered_by']}")
            run3.font.size = Pt(8)
            run3.font.color.rgb = GRAY

    doc.add_page_break()

    # ════════════════════════════════════════
    # FRAMEWORK ALIGNMENT
    # ════════════════════════════════════════
    section_num = 5
    for fw_id, fw_info in selected_fw.items():
        _heading(f"{section_num}. {fw_info['name']} Alignment")
        _para(f"{fw_info['full']} ({fw_info['version']})", color=GRAY, size=9)

        rows = []
        for ctrl in assessment["controls"]:
            ref = ctrl["mappings"].get(fw_id, "")
            if ref and ref != "—":
                rows.append([ref, f"{ctrl['id']} — {ctrl['title']}", ctrl["status"]])
        if rows:
            _table([fw_info["name"], "Control", "Status"], rows)
        else:
            _para("No controls mapped to this framework.", color=GRAY, size=9)

        section_num += 1
        doc.add_page_break()

    # ════════════════════════════════════════
    # RECOMMENDATIONS
    # ════════════════════════════════════════
    _heading(f"{section_num}. Prioritized Recommendations")
    for rec in recs:
        pr = "HIGH" if rec["weight"] >= 8 else "MEDIUM" if rec["weight"] >= 6 else "LOW"
        pc = RED if pr == "HIGH" else ORANGE if pr == "MEDIUM" else RGBColor(0xCA, 0x8A, 0x04)
        p = doc.add_paragraph()
        run = p.add_run(f"[{pr}] ")
        run.bold = True
        run.font.color.rgb = pc
        run = p.add_run(f"{rec['id']}: {rec['title']}")
        run.bold = True
        doc.add_paragraph(rec["recommendation"]).runs[0].font.size = Pt(9)
        fw_refs = ", ".join(f"{FRAMEWORKS[k]['name']} {v}" for k, v in rec["mappings"].items() if v != "—")
        if fw_refs:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(fw_refs)
            run2.italic = True
            run2.font.size = Pt(8)
            run2.font.color.rgb = GRAY

    if not recs:
        _para("All controls passed — full compliance achieved.", bold=True, color=GREEN)

    section_num += 1

    # ════════════════════════════════════════
    # GRC SUMMARY (DOCX)
    # ════════════════════════════════════════
    grc = assessment.get("grc")
    if grc and not grc.get("error"):
        doc.add_page_break()
        _heading(f"{section_num}. Governance, Risk & Compliance Overview")

        rk = grc.get("risks", {})
        if rk.get("total", 0) > 0:
            _heading("Risk Register", level=2)
            doc.add_paragraph(f"Total: {rk['total']} risks — Open: {rk.get('open',0)} — Critical: {rk.get('critical',0)}")
            top5 = rk.get("top5", [])
            if top5:
                _table(["Risk", "Score", "Status"],
                       [[r["title"], str(r["score"]), r["status"]] for r in top5])

        ast = grc.get("assets", {})
        if ast.get("total", 0) > 0:
            _heading("Asset Inventory", level=2)
            by_type = ast.get("by_type", {})
            _table(["Type", "Count"],
                   [[t.capitalize(), str(c)] for t, c in sorted(by_type.items())])

        pol = grc.get("policies", {})
        if pol.get("total", 0) > 0:
            _heading("Policy Status", level=2)
            doc.add_paragraph(f"{pol['total']} policies: {pol.get('approved',0)} approved, {pol.get('draft',0)} draft")

        vnd = grc.get("vendors", {})
        if vnd.get("total", 0) > 0:
            _heading("Third-Party Risk", level=2)
            doc.add_paragraph(f"{vnd['total']} vendors: {vnd.get('high_risk',0)} high/critical risk")

        af = grc.get("audit_findings_open", 0)
        if af > 0:
            _heading("Open Audit Findings", level=2)
            p = doc.add_paragraph()
            run = p.add_run(f"{af} finding(s) require remediation")
            run.font.color.rgb = ORANGE

        trend = grc.get("compliance_trend", {})
        if any(v is not None for v in trend.values()):
            _heading("Compliance Trend", level=2)
            labels = {"m1": "vs. 1 month ago", "m3": "vs. 3 months ago",
                      "m6": "vs. 6 months ago", "m12": "vs. 12 months ago"}
            rows = []
            for k, label in labels.items():
                v = trend.get(k)
                if v is not None:
                    rows.append([label, f"{'+' if v>0 else ''}{v} points"])
            if rows:
                _table(["Period", "Delta"], rows)

        section_num += 1

    doc.add_page_break()

    # ════════════════════════════════════════
    # ATTESTATION
    # ════════════════════════════════════════
    _heading("COMPLIANCE ATTESTATION")
    org_name = company if company else "the assessed organization"
    doc.add_paragraph(
        f"This report certifies that {org_name} has undergone an automated "
        f"cybersecurity compliance assessment on {now.strftime('%B %d, %Y')}, "
        f"covering {total} security controls aligned with {fw_names}.")
    doc.add_paragraph(
        f"The overall compliance score is {score}/100, corresponding to a "
        f"grade of {grade}. {passed} out of {total} controls passed.")
    doc.add_paragraph()

    info2 = []
    if company:
        info2.append(("Organization", company))
    info2 += [
        ("Assessment Date", now.strftime("%Y-%m-%d %H:%M:%S")),
        ("Report ID", report_id),
        ("Frameworks", fw_names),
        ("Method", "Automated continuous monitoring & control verification"),
    ]
    t2 = doc.add_table(rows=len(info2), cols=2)
    t2.style = 'Light List'
    for i, (label, value) in enumerate(info2):
        t2.rows[i].cells[0].text = label
        t2.rows[i].cells[1].text = value
        for run in t2.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()
    if fully_compliant:
        _para("All controls are in compliance. This document serves as a certificate of cybersecurity compliance.",
              color=GREEN, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _para("This assessment reflects the state at the time of generation. "
              "Self-assessed controls have not been independently verified.",
              color=GRAY, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save to bytes
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()
    docx_buf.close()
    return docx_bytes
